#!/usr/bin/env python3
"""
Document Conversion Pipeline

This module implements Section 1.2 of the strategic framework: "Pre-processing: Batch 
Conversion of DOCX to TXT". It provides automated batch conversion of Microsoft Word 
documents to plain text format for ingestion by the knowledge graph engine.

Key Features:
- Batch processing of entire directories
- Multiple conversion libraries with fallback
- Parallel processing for performance
- Comprehensive error handling and logging
- Format preservation options

Author: AI Mapping Knowledge Graph System
"""

import os
import logging
import multiprocessing
from pathlib import Path
from typing import List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

# Document conversion libraries
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentConverter:
    """
    Batch document converter following the framework's recommendations.
    
    Supports multiple conversion methods with automatic fallback:
    1. docx2txt (recommended for simplicity)
    2. python-docx (for granular control)
    3. pypandoc (for maximum compatibility)
    """
    
    def __init__(self, preserve_formatting: bool = False):
        """
        Initialize the document converter.
        
        Args:
            preserve_formatting: Whether to attempt to preserve basic formatting
        """
        self.preserve_formatting = preserve_formatting
        self.conversion_stats = {
            'total_files': 0,
            'successful': 0,
            'failed': 0,
            'errors': []
        }
        
        # Check available conversion methods
        self.available_methods = []
        if DOCX2TXT_AVAILABLE:
            self.available_methods.append('docx2txt')
        if PYTHON_DOCX_AVAILABLE:
            self.available_methods.append('python-docx')
        if PYPANDOC_AVAILABLE:
            self.available_methods.append('pypandoc')
        
        if not self.available_methods:
            raise RuntimeError("No document conversion libraries available. Please install docx2txt, python-docx, or pypandoc")
        
        logger.info(f"Available conversion methods: {', '.join(self.available_methods)}")
    
    def convert_single_document(self, docx_path: Path, output_path: Path, 
                              method: Optional[str] = None) -> bool:
        """
        Convert a single DOCX file to TXT.
        
        Args:
            docx_path: Path to input DOCX file
            output_path: Path for output TXT file
            method: Specific conversion method to use (optional)
            
        Returns:
            True if conversion successful, False otherwise
        """
        if not docx_path.exists():
            logger.error(f"Input file does not exist: {docx_path}")
            return False
        
        if not docx_path.suffix.lower() == '.docx':
            logger.warning(f"File does not appear to be DOCX: {docx_path}")
        
        # Create output directory if it doesn't exist
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Determine conversion method
        methods_to_try = [method] if method else self.available_methods
        
        for conv_method in methods_to_try:
            if conv_method not in self.available_methods:
                continue
                
            try:
                if conv_method == 'docx2txt':
                    text = self._convert_with_docx2txt(docx_path)
                elif conv_method == 'python-docx':
                    text = self._convert_with_python_docx(docx_path)
                elif conv_method == 'pypandoc':
                    text = self._convert_with_pypandoc(docx_path)
                else:
                    continue
                
                # Write to output file
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                logger.debug(f"Converted {docx_path.name} using {conv_method}")
                return True
                
            except Exception as e:
                logger.warning(f"Conversion failed with {conv_method} for {docx_path.name}: {e}")
                continue
        
        logger.error(f"All conversion methods failed for {docx_path.name}")
        return False
    
    def _convert_with_docx2txt(self, docx_path: Path) -> str:
        """Convert using docx2txt library (recommended method)"""
        text = docx2txt.process(str(docx_path))
        return self._clean_text(text)
    
    def _convert_with_python_docx(self, docx_path: Path) -> str:
        """Convert using python-docx library (provides more control)"""
        doc = Document(str(docx_path))
        
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        text = "\n\n".join(paragraphs)
        return self._clean_text(text)
    
    def _convert_with_pypandoc(self, docx_path: Path) -> str:
        """Convert using pypandoc (most compatible)"""
        text = pypandoc.convert_file(str(docx_path), 'plain')
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Basic text cleaning
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Collapse multiple blank lines
        text = re.sub(r'[ \t]+', ' ', text)  # Collapse multiple spaces/tabs
        
        # Remove or replace special characters that might cause issues
        text = text.replace('\r', '\n')  # Normalize line endings
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)  # Remove control chars
        
        # Trim and ensure proper ending
        text = text.strip()
        if text and not text.endswith('\n'):
            text += '\n'
        
        return text
    
    def batch_convert(self, input_dir: str, output_dir: str, 
                     max_workers: Optional[int] = None, 
                     pattern: str = "*.docx") -> Tuple[int, int]:
        """
        Batch convert all DOCX files in a directory.
        
        Args:
            input_dir: Directory containing DOCX files
            output_dir: Directory for output TXT files
            max_workers: Number of parallel workers (default: CPU count)
            pattern: File pattern to match (default: "*.docx")
            
        Returns:
            Tuple of (successful_conversions, failed_conversions)
        """
        input_path = Path(input_dir)
        output_path = Path(output_dir)
        
        if not input_path.exists():
            logger.error(f"Input directory does not exist: {input_dir}")
            return 0, 0
        
        # Find all DOCX files
        docx_files = list(input_path.rglob(pattern))
        if not docx_files:
            logger.warning(f"No files matching {pattern} found in {input_dir}")
            return 0, 0
        
        logger.info(f"Found {len(docx_files)} files to convert")
        
        # Reset stats
        self.conversion_stats = {
            'total_files': len(docx_files),
            'successful': 0,
            'failed': 0,
            'errors': []
        }
        
        # Determine optimal number of workers
        if max_workers is None:
            max_workers = min(multiprocessing.cpu_count(), len(docx_files))
        
        logger.info(f"Starting batch conversion with {max_workers} workers")
        start_time = time.time()
        
        # Parallel processing
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all conversion tasks
            future_to_file = {}
            for docx_file in docx_files:
                # Preserve directory structure in output
                relative_path = docx_file.relative_to(input_path)
                txt_file = output_path / relative_path.with_suffix('.txt')
                
                future = executor.submit(self.convert_single_document, docx_file, txt_file)
                future_to_file[future] = (docx_file, txt_file)
            
            # Collect results
            for future in as_completed(future_to_file):
                docx_file, txt_file = future_to_file[future]
                try:
                    success = future.result()
                    if success:
                        self.conversion_stats['successful'] += 1
                    else:
                        self.conversion_stats['failed'] += 1
                        self.conversion_stats['errors'].append(str(docx_file))
                        
                except Exception as e:
                    self.conversion_stats['failed'] += 1
                    self.conversion_stats['errors'].append(f"{docx_file}: {str(e)}")
                    logger.error(f"Conversion error for {docx_file}: {e}")
                
                # Progress reporting
                completed = self.conversion_stats['successful'] + self.conversion_stats['failed']
                if completed % 10 == 0 or completed == len(docx_files):
                    progress = (completed / len(docx_files)) * 100
                    logger.info(f"Progress: {completed}/{len(docx_files)} ({progress:.1f}%)")
        
        elapsed_time = time.time() - start_time
        
        # Final report
        logger.info("=== BATCH CONVERSION COMPLETE ===")
        logger.info(f"Total files: {self.conversion_stats['total_files']}")
        logger.info(f"Successful: {self.conversion_stats['successful']}")
        logger.info(f"Failed: {self.conversion_stats['failed']}")
        logger.info(f"Success rate: {(self.conversion_stats['successful'] / self.conversion_stats['total_files']) * 100:.1f}%")
        logger.info(f"Time elapsed: {elapsed_time:.2f} seconds")
        
        if self.conversion_stats['errors']:
            logger.warning("Failed conversions:")
            for error in self.conversion_stats['errors'][:10]:  # Show first 10 errors
                logger.warning(f"  {error}")
            if len(self.conversion_stats['errors']) > 10:
                logger.warning(f"  ... and {len(self.conversion_stats['errors']) - 10} more")
        
        return self.conversion_stats['successful'], self.conversion_stats['failed']
    
    def get_conversion_stats(self) -> dict:
        """Get detailed conversion statistics"""
        return self.conversion_stats.copy()

def main():
    """Main execution function for standalone use"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Batch convert DOCX files to TXT")
    parser.add_argument("input_dir", help="Directory containing DOCX files")
    parser.add_argument("output_dir", help="Directory for output TXT files")
    parser.add_argument("--workers", type=int, help="Number of parallel workers")
    parser.add_argument("--pattern", default="*.docx", help="File pattern to match")
    parser.add_argument("--preserve-formatting", action="store_true", 
                       help="Attempt to preserve basic formatting")
    
    args = parser.parse_args()
    
    # Create converter
    converter = DocumentConverter(preserve_formatting=args.preserve_formatting)
    
    # Run batch conversion
    successful, failed = converter.batch_convert(
        args.input_dir, 
        args.output_dir, 
        max_workers=args.workers,
        pattern=args.pattern
    )
    
    # Exit with appropriate code
    if failed == 0:
        logger.info("All conversions successful!")
        exit(0)
    elif successful > 0:
        logger.warning(f"Partial success: {successful} succeeded, {failed} failed")
        exit(1)
    else:
        logger.error("All conversions failed!")
        exit(2)

if __name__ == "__main__":
    main() 